"""
Detailed expense-management analysis for FamilyHub, sourced from the
Quach Minh Phat / MoMo user interview (31/07/2026).
Run: /opt/homebrew/bin/python3 gen_expense_management_analysis.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# FamilyHub design tokens (DESIGN.md) — semantic, no colored fills, border-led hierarchy
BRAND      = RGBColor(0x2E, 0x9E, 0x6B)
BRAND_INK  = RGBColor(0x1F, 0x7E, 0x52)
INK        = RGBColor(0x19, 0x10, 0x22)
INK_2      = RGBColor(0x3A, 0x2F, 0x45)
MUTED      = RGBColor(0x8A, 0x84, 0x94)
GOOD       = RGBColor(0x1A, 0x9D, 0x5F)
DANGER     = RGBColor(0xE0, 0x32, 0x2F)
AMBER      = RGBColor(0xB8, 0x73, 0x0B)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
HAIRLINE   = "E9E4EE"

FONT = "Inter"


def set_run(run, *, font=FONT, size=11, bold=False, italic=False, color=INK_2):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rFonts.set(qn(f"w:{attr}"), font)


def add_paragraph(doc, text="", *, size=11, bold=False, italic=False,
                  color=INK_2, align=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before=0, space_after=4, line_spacing=1.35):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if text:
        run = p.add_run(text)
        set_run(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(22)
    pf.space_after = Pt(8)
    run = p.add_run(text)
    set_run(run, size=19, bold=True, color=BRAND_INK)
    _add_bottom_border(p, BRAND, sz=12)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(14)
    pf.space_after = Pt(4)
    run = p.add_run(text)
    set_run(run, size=13.5, bold=True, color=INK)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(9)
    pf.space_after = Pt(2)
    run = p.add_run(text)
    set_run(run, size=11.5, bold=True, color=BRAND_INK)
    return p


def add_bullet(doc, text, *, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.3
    if bold_lead:
        r = p.add_run(bold_lead)
        set_run(r, size=11, bold=True, color=INK)
        r2 = p.add_run(text)
        set_run(r2, size=11, color=INK_2)
    else:
        r = p.add_run(text)
        set_run(r, size=11, color=INK_2)
    return p


def add_quote(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(0.6)
    pf.space_before = Pt(4)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.3
    r = p.add_run('"' + text + '"')
    set_run(r, size=10.5, italic=True, color=MUTED)
    return p


def add_implication(doc, text):
    """Left-accent callout, white fill only (FamilyHub: no colored card fills)."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    _set_cell_bg(cell, "FFFFFF")
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, csz, ccolor in (
        ("top", "4", HAIRLINE),
        ("right", "4", HAIRLINE),
        ("bottom", "4", HAIRLINE),
        ("left", "24", "2E9E6B"),
    ):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), csz)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), ccolor)
        tcBorders.append(b)
    tcPr.append(tcBorders)

    cell.text = ""
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run("Ý NGHĨA CHO FAMILYHUB")
    set_run(r1, size=9, bold=True, color=BRAND_INK)

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    p2.paragraph_format.line_spacing = 1.3
    r2 = p2.add_run(text)
    set_run(r2, size=10.5, color=INK_2)

    add_paragraph(doc, "", space_after=6)


def _add_bottom_border(paragraph, color: RGBColor, sz=8):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, color=HAIRLINE, sz="6"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _set_cell_text(cell, text, *, bold=False, color=INK, size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.25
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc, headers, rows, *, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = w

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_bg(cell, "2E9E6B")
        _set_cell_border(cell, color="2E9E6B")
        _set_cell_text(cell, h, bold=True, color=WHITE, size=10)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            _set_cell_bg(cell, "FFFFFF")
            _set_cell_border(cell, color=HAIRLINE)
            _set_cell_text(cell, val, size=10, color=INK_2)

    return table


def build():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(11)

    # ==================== COVER ====================
    add_paragraph(doc, "FAMILYHUB • RESEARCH INPUT",
                  size=10, bold=True, color=BRAND_INK, space_before=8, space_after=4)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Phân tích chuyên sâu: Quản lý chi tiêu")
    set_run(r, size=23, bold=True, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Rút ra từ phỏng vấn Quách Minh Phát (MoMo, 31/07/2026) — đối chiếu với các module chi tiêu của FamilyHub")
    set_run(r, size=13, italic=True, color=MUTED)
    _add_bottom_border(p, BRAND, sz=18)

    meta = [
        ("Nguồn dữ liệu", "Phỏng vấn người dùng MoMo — Quách Minh Phát, 23 tuổi, ~103 phút, 31/07/2026"),
        ("Phạm vi tài liệu này", "Chỉ trích xuất và phân tích các đoạn liên quan tới quản lý chi tiêu "
                                 "(ngân sách, chia sẻ chi tiêu, hoàn tiền, phân loại, tự động hoá, tín dụng)"),
        ("Mục đích", "Làm đầu vào thiết kế cho các module chi tiêu của FamilyHub "
                     "(budget, transactions, expense-detail, requests, expense-photos/OCR)"),
        ("Bối cảnh người được phỏng vấn", "Sống chung với partner (không phải quan hệ cha mẹ – con); "
                                          "tự xây một hệ thống quản lý chi tiêu chung thủ công bằng Google Sheets "
                                          "vì các công cụ có sẵn (Momo, ngân hàng) không đáp ứng được nhu cầu"),
    ]
    t = doc.add_table(rows=len(meta), cols=2)
    t.autofit = False
    for i, (k, v) in enumerate(meta):
        c1, c2 = t.rows[i].cells
        c1.width = Cm(4.4)
        c2.width = Cm(11.2)
        _set_cell_bg(c1, "FFFFFF")
        _set_cell_bg(c2, "FFFFFF")
        _set_cell_border(c1, HAIRLINE)
        _set_cell_border(c2, HAIRLINE)
        _set_cell_text(c1, k, bold=True, color=BRAND_INK, size=10)
        _set_cell_text(c2, v, size=10, color=INK_2)

    add_paragraph(doc, "", space_after=8)

    # ==================== 1. TÓM TẮT ĐIỀU HÀNH ====================
    add_h1(doc, "1. Tóm tắt điều hành")

    add_paragraph(
        doc,
        "Quách Minh Phát và partner (sống chung, không kết hôn) đã tự xây một \"FamilyHub thu nhỏ\" bằng tay: "
        "ngân sách cố định theo tháng, một người giữ vai trò \"mastermind\" tổng hợp chi tiêu, ứng tiền trước và hoàn lại sau, "
        "tất cả ghi vào một file Google Sheets 2 cột (ngày — số tiền). Đây không phải vì họ không biết đến tính năng quản lý chi tiêu "
        "của Momo — họ biết, và thậm chí đánh giá cao tính năng phân loại giao dịch — mà vì hai giới hạn cụ thể của các công cụ hiện có: "
        "(1) không chia sẻ được dữ liệu giao dịch cá nhân giữa hai người mà không phải gộp tiền vào một quỹ chung, và "
        "(2) không phân biệt được một giao dịch nào là \"chung\", \"riêng\", hay \"chung một phần\".",
        space_after=6,
    )
    add_paragraph(
        doc,
        "Đây chính là khoảng trống mà FamilyHub — một sản phẩm được thiết kế từ đầu cho nhiều thành viên, không phải một ví cá nhân "
        "gắn thêm tính năng chia sẻ — có thể giải quyết triệt để hơn Momo. Tài liệu này đi sâu vào bảy khía cạnh cụ thể của cách "
        "Phát quản lý chi tiêu, mỗi khía cạnh kèm theo hàm ý thiết kế cụ thể cho FamilyHub.",
        space_after=4,
    )

    # ==================== 2. PHÂN TÍCH CHUYÊN SÂU ====================
    add_h1(doc, "2. Phân tích chuyên sâu theo từng khía cạnh")

    # 2.1
    add_h2(doc, "2.1. Kiến trúc ngân sách: khoản cố định tách khỏi khoản biến động, một người giữ vai trò \"mastermind\"")
    add_paragraph(
        doc,
        "Phát và partner chia ngân sách tháng thành hai loại rõ rệt: khoản cố định (tiền nhà — không đổi trong suốt hợp đồng thuê) "
        "và khoản biến động có trần (ngân sách ăn uống — mỗi người góp một phần cố định vào đầu tháng). Ranh giới ngân sách ăn uống "
        "được đặt ra chủ động để tự giới hạn hành vi phát sinh (\"ngay cả khi hôm nào đó cao hứng đi ăn ở chỗ khác thì nó vẫn nằm trong "
        "budget đó\") — tức là ngân sách đóng vai trò một hàng rào tâm lý, không chỉ một con số kế toán. Một người (partner của Phát, "
        "một kiến trúc sư) tình nguyện nhận vai trò điều phối trung tâm — theo dõi từng ngày, tổng hợp, và duy trì \"nguồn sự thật\" duy nhất.",
        space_after=4,
    )
    add_quote(doc, "Tụi em sẽ xích với nhau là trong cái tháng đó tụi em sẽ có những khoản fix. Ví dụ tiền nhà là fix với nhau rồi... tiền ăn tụi em nhận ra được là cần phải có một cái budget, để nó tránh cái việc mình exit.")
    add_quote(doc, "Partner em là cái người mastermind, là cái người quản lý những cái chuyện đó — bạn sẽ track xem ngày hôm nay ăn cái gì.")
    add_implication(
        doc,
        "FamilyHub nên hỗ trợ rõ hai lớp ngân sách trong cùng một gia đình: khoản cố định (rent-like, không cần theo dõi sát) và khoản "
        "biến động có trần theo kỳ (food-like, cần cảnh báo khi gần vượt). Nên cho phép chỉ định một thành viên làm \"người điều phối\" "
        "cho một hạng mục ngân sách cụ thể — không nhất thiết phải là chủ hộ/người tạo family — để phản ánh đúng cách các cặp/hộ thực sự "
        "phân công việc quản lý tiền trong đời thực.",
    )

    # 2.2
    add_h2(doc, "2.2. Vấn đề cốt lõi: một giao dịch có thể vừa là \"của tôi\" vừa là \"của chung\" — nhị phân chung/riêng là không đủ")
    add_paragraph(
        doc,
        "Đây là insight quan trọng nhất của toàn bộ đoạn phỏng vấn về chi tiêu. Khi được hỏi bao nhiêu phần trăm chi tiêu ăn uống hằng tháng "
        "là \"chung\", Phát trả lời rằng ranh giới rất khó phân định — cùng một khoản ăn uống có thể là hoàn toàn cá nhân (tự trả khi ăn "
        "một mình hoặc ăn với bạn ngoài), hoàn toàn chung (bữa ăn của cả hai), hoặc chung một phần (nhóm 3-4 người ăn chung, một người đứng "
        "ra trả, rồi tách phần của từng người). Vấn đề kỹ thuật cụ thể: khi dữ liệu giao dịch đổ về từ ngân hàng/ví, nó chỉ là MỘT con số "
        "tổng — không có cách nào tách được \"bao nhiêu trong đó là của tôi, bao nhiêu là của người khác\" ở cấp giao dịch.",
        space_after=4,
    )
    add_quote(doc, "Em nghĩ là nó là có cả hai. Nó rất là khó để có thể phân định được... khi mà dữ liệu nó đổ về thì mình sẽ không có tách ra được là cái nào của ai của ai.")
    add_quote(doc, "Em hình dung nó sẽ có những cái trường dữ liệu mà mình có thể tách ra được... có thể nó cụ thể hóa hơn nữa là trong cái cục 72.000 đó chẳng hạn thì bao nhiêu tiền của tôi, bao nhiêu tiền của người kia kiểu vậy.")
    add_implication(
        doc,
        "Đây là yêu cầu về mô hình dữ liệu, không phải giao diện: mỗi giao dịch cần hỗ trợ SPLIT ở cấp bản ghi (không chỉ gắn nhãn "
        "\"chung\"/\"riêng\" nhị phân), với các phần chia có thể gán cho từng thành viên — kể cả người không phải chủ giao dịch. "
        "Đây chính là cơ chế \"split the bill\" kinh điển (như Splitwise), nhưng cần tích hợp thẳng vào transaction record của FamilyHub "
        "thay vì là một tính năng phụ. Nên xem lại `61-expense-detail.js` / `60-transactions.js` để xác nhận đã có trường "
        "split-by-member hay chưa, và nếu chưa, đây nên là ưu tiên cao vì nó là nút thắt khiến người dùng thực tế (Phát) quay lưng với "
        "công cụ có sẵn để dùng Google Sheets.",
    )

    # 2.3
    add_h2(doc, "2.3. Ứng trước và hoàn tiền: gia đình vận hành như một hệ thống tín dụng vi mô nội bộ")
    add_paragraph(
        doc,
        "Cơ chế cụ thể: đầu tháng Phát gửi trước phần ngân sách ăn uống của mình cho partner; partner dùng số dư đó ứng trước khi có chi "
        "tiêu phát sinh; cuối ngày hoặc cuối tuần, partner hoàn lại phần chênh lệch (refund) nếu chi tiêu thực tế thấp hơn phần đã ứng, "
        "hoặc yêu cầu bù thêm nếu cao hơn. Đây bản chất là một vòng lặp ứng trước → ghi nhận → đối soát → hoàn tiền, lặp lại theo chu kỳ "
        "ngày/tuần chứ không phải theo giao dịch đơn lẻ.",
        space_after=4,
    )
    add_quote(doc, "Đầu tháng là em sẽ gửi trước cái phần budget tiền ăn đó cho bạn, cho nên là đến cuối ngày hoặc là đến cuối tuần bạn sẽ refund cái số tiền đó về cho em tại vì em đang phải lấy một cái phần tiền khác để trả trước.")
    add_implication(
        doc,
        "FamilyHub's `64-requests.js` hiện đang xử lý luồng \"request/review\" mang tính phê duyệt (reaction-based) — khác với nhu cầu ở đây, "
        "vốn là một sổ cái công nợ nội bộ (running balance) giữa 2 thành viên: ai đang ứng hộ ai bao nhiêu, tự động cấn trừ theo thời gian. "
        "Nên cân nhắc một \"IOU ledger\" giữa các cặp thành viên — không cần thanh toán thật qua ngân hàng, chỉ cần một con số công nợ "
        "hiển thị rõ và có thể \"đánh dấu đã tất toán\" theo chu kỳ do gia đình tự chọn (ngày/tuần/tháng).",
    )

    # 2.4
    add_h2(doc, "2.4. Rào cản của \"quỹ nhóm\": ép người dùng chọn giữa gộp tiền hoặc không thấy gì cả")
    add_paragraph(
        doc,
        "Phát hiểu rõ cơ chế quỹ nhóm của Momo — có thể yêu cầu (request) từ quỹ nếu là thành viên, giao dịch được thống kê ở cấp quỹ ngay "
        "cả khi thanh toán từ nguồn ngoài. Nhưng anh chủ động không dùng nó cho việc theo dõi chi tiêu ăn uống với partner, vì nhu cầu thực "
        "sự không phải là quản lý MỘT nguồn tiền chung — mà là XEM ĐƯỢC chi tiêu CÁ NHÂN của nhau. Quỹ nhóm buộc người dùng phải di chuyển "
        "tiền vào một pool trước khi có thể theo dõi nó cùng nhau; Phát muốn giữ tiền trong tài khoản riêng của mỗi người nhưng vẫn nhìn thấy "
        "nhật ký chi tiêu của đối phương theo thời gian thực.",
        space_after=4,
    )
    add_quote(doc, "Quỹ nhóm thì mình biết được cái chuyện là tiền ra tiền vô như thế nào... Nhưng mà giả sử những cái transaction đó mình có thể share được với nhau hay không thì em không sure là mình đã có tính năng đó hay chưa.")
    add_quote(doc, "Nếu như mà nó automatic được đến như vậy và em có thể coi, cả em với partner của em coi được, share chung cái data đó với nhau, thì likely là bạn em sẽ nhập lên trên Momo thay vì làm cái Google Sheets thì nó quá cực.")
    add_implication(
        doc,
        "Đây là điểm khác biệt chiến lược rõ nhất mà FamilyHub có thể tận dụng so với ví điện tử: một chế độ \"CHIA SẺ TẦM NHÌN, KHÔNG "
        "CHIA SẺ TIỀN\" — cho phép một thành viên cấp quyền xem (toàn bộ hoặc theo danh mục) nhật ký chi tiêu cá nhân của mình cho thành "
        "viên khác trong gia đình, độc lập với việc có quỹ chung hay không. Quỹ nhóm/quỹ chung vẫn nên tồn tại song song cho trường hợp cần "
        "gộp tiền thật (ví dụ quỹ du lịch), nhưng không nên là điều kiện bắt buộc để có được visibility.",
    )

    # 2.5
    add_h2(doc, "2.5. Dữ liệu chi tiêu luôn bị \"gãy\" khi hộ gia đình dùng nhiều app thanh toán song song")
    add_paragraph(
        doc,
        "Vì chuyển đổi linh hoạt giữa Momo và app ngân hàng (khi app này lỗi thì dùng app kia), báo cáo chi tiêu cuối tháng của Phát trên "
        "Momo không phản ánh đầy đủ thực tế — tính năng đồng bộ ngân hàng hiện tại chỉ lấy số dư tổng, không lấy chi tiết từng giao dịch. "
        "Kết quả là để có dữ liệu đầy đủ, gia đình vẫn phải quay lại nhập tay từ nhiều nguồn (ngân hàng, ví điện tử, hoá đơn giấy) vào một "
        "nơi duy nhất — đúng bài toán mà công cụ chụp hoá đơn/OCR đang được phác thảo trong `reference/ocr_expense_logging_flow*.drawio.xml` "
        "của dự án hướng tới giải quyết.",
        space_after=4,
    )
    add_quote(doc, "Không phải cái giao dịch nào của em nó cũng được đổ tới Momo... dữ liệu mà em có được vào cuối tháng nó sẽ không reflect được.")
    add_quote(doc, "Tụi em đi quá nhiều, tụi em mua quá nhiều dữ liệu để tỏa tẻ. Quẹt rất là nhiều lần, tụi em không thể nào cầm hết 10 cái bill đó trong tay được. Thì buộc phải lấy dữ liệu từ either Momo hoặc là từ ngân hàng.")
    add_implication(
        doc,
        "Xác nhận rằng luồng OCR/chụp hoá đơn mà FamilyHub đang thiết kế là đúng hướng — nhưng nên mở rộng ngoài use-case \"chụp hoá đơn giấy\" "
        "sang cả \"chụp màn hình lịch sử giao dịch ngân hàng/ví\" như một nguồn nhập liệu thay thế, vì đây chính xác là hành vi thực tế mà "
        "Phát và partner đang làm thủ công (soi lại lịch sử giao dịch Momo/Vietcombank để tổng hợp).",
    )

    # 2.6
    add_h2(doc, "2.6. Tâm lý ngân sách: cảnh báo phải trả lời \"tôi còn lại bao nhiêu\", không phải \"tôi tiêu bao nhiêu\"")
    add_paragraph(
        doc,
        "Phát chỉ ra một cơ chế tâm lý cụ thể và có thể thiết kế được: người dùng hoảng (\"phát hoảng\") không phải vì tổng số tiền sẽ bị "
        "trừ trong tháng, mà vì con số dư khả dụng giảm đột ngột và không được giải thích trước. Giải pháp anh tự đề xuất — độc lập, "
        "không được gợi ý trước — là một thông báo đầu tháng liệt kê cụ thể các khoản dự kiến bị trừ (điện, nước, subscription) dựa trên "
        "lịch sử, cùng với số dư dự kiến còn lại sau khi trừ hết. Anh cũng nhắc tới phương pháp \"5 hũ\" của người Do Thái như một mô hình "
        "tham chiếu: tiền được thấy rõ đang \"chảy vào\" từng mục đích cụ thể, thay vì phải truy lại từng giao dịch để đoán tiền đã đi đâu.",
        space_after=4,
    )
    add_quote(doc, "Mọi người sẽ phát hoảng không phải là vì tổng cái số tiền mình phải trả, mà mọi người sẽ phát hoảng vì cái số dư còn lại mà nó hiện ra.")
    add_quote(doc, "Người Do Thái dạy tiền có 5 hũ... người ta sẽ nhìn thấy được là cái tiền từ trong cái hũ đấy nó rơi ra, không phải là như em sẽ phải nhìn lại hết mọi transaction mới đoán ra là cái tiền bị mất là do từ đâu.")
    add_quote(doc, "Đầu tháng em nhận được một cái notice... nó sẽ list ra rất là cụ thể là dựa trên lịch sử mà bạn đã trả bill... thì sẽ có một cái hàng là tổng số tiền mà bạn likely sẽ trả, và số còn dư là của bạn.")
    add_implication(
        doc,
        "FamilyHub's `20-budget.js` nên ưu tiên hiển thị \"số dư khả dụng sau các khoản cố định sắp tới\" như một con số chính, không chỉ "
        "\"đã chi bao nhiêu / còn lại bao nhiêu trong ngân sách\". Envelope/bucket budgeting (tương đương \"5 hũ\") là mô hình tâm lý người "
        "dùng tự nhắc tới không cần gợi ý — nên cân nhắc cho phép mỗi thành viên tạo các \"hũ\" mục đích riêng (không chỉ category chi tiêu) "
        "và thấy tiền \"rơi vào\" hũ theo thời gian thực khi lương về.",
    )

    # 2.7
    add_h2(doc, "2.7. Rủi ro tài chính ẩn giữa các thành viên: một người dùng ví trả sau âm thầm có thể tạo khủng hoảng cho cả hai")
    add_paragraph(
        doc,
        "Phát kể lại một tình huống thực tế trong vòng quan hệ gần: bạn của bạn (bồ của người bạn cùng nhà) mượn điện thoại của partner "
        "mình, dùng ví trả sau chi tiêu khắp nơi, và chỉ đến cuối tháng partner mới phát hiện một khoản nợ lớn — người phải đứng ra trả "
        "là partner, không phải người tiêu. Đây là ví dụ cụ thể cho thấy trong một hộ/cặp chia sẻ thiết bị hoặc tài khoản, rủi ro tín dụng "
        "của một người có thể lặng lẽ trở thành gánh nặng tài chính của người khác nếu không có cảnh báo sớm.",
        space_after=4,
    )
    add_quote(doc, "Họ cầm cái máy của của bồ mình và họ dùng ví trả sau và họ quẹt đi đủ chỗ hết trơn á, cho đến cuối tháng thì partner mới phát hiện ra là có một cục nợ lớn như vậy.")
    add_quote(doc, "Cái trap đó nó rất là khó để đi ra được, đặc biệt là khi mình không có safety net ở trong người.")
    add_implication(
        doc,
        "Với FamilyHub, đây là lập luận ủng hộ hiển thị nợ/khoản ứng trước (mục 2.3) và hạn mức chi tiêu ở cấp thành viên có thể nhìn thấy "
        "được bởi người điều phối ngân sách gia đình (không cần chặn giao dịch, chỉ cần cảnh báo sớm) — đặc biệt quan trọng khi FamilyHub "
        "mở rộng ra dùng cho các cặp cha mẹ – con hoặc gia đình có thành viên trẻ mới bắt đầu tự chủ tài chính, nơi rủi ro \"nợ ẩn\" cao hơn.",
    )

    # ==================== 3. ĐỐI CHIẾU VỚI FAMILYHUB HIỆN TẠI ====================
    add_h1(doc, "3. Bảng đối chiếu: Insight → Module liên quan → Đề xuất")

    mapping = [
        ("Split chung/riêng ở cấp giao dịch (2.2)", "60-transactions.js, 61-expense-detail.js",
         "Thêm trường chia theo % / số tiền cho từng thành viên trên một giao dịch, không chỉ nhãn chung/riêng"),
        ("Ứng trước & hoàn tiền theo chu kỳ (2.3)", "64-requests.js",
         "Bổ sung sổ công nợ nội bộ (running IOU) giữa 2 thành viên, tách khỏi luồng request/review hiện tại"),
        ("Chia sẻ tầm nhìn không cần gộp quỹ (2.4)", "20-budget.js (mới)",
         "Chế độ cấp quyền xem nhật ký chi tiêu cá nhân cho thành viên khác, độc lập với quỹ chung"),
        ("Dữ liệu gãy do đa nguồn thanh toán (2.5)", "55-expense-photos-writes.js, OCR flow (reference/)",
         "Mở rộng OCR sang chụp lịch sử giao dịch ngân hàng/ví, không chỉ hoá đơn giấy"),
        ("Cảnh báo số dư khả dụng, không phải tổng chi (2.6)", "20-budget.js, 22-home.js",
         "Ưu tiên hiển thị số dư dự kiến sau các khoản cố định sắp tới; cân nhắc mô hình \"hũ\" (envelope)"),
        ("Nợ ẩn giữa thành viên (2.7)", "64-requests.js, 20-budget.js",
         "Cảnh báo sớm cho người điều phối khi một thành viên có khoản ứng/nợ vượt ngưỡng"),
    ]
    add_table(
        doc,
        ["Insight", "Module liên quan", "Đề xuất"],
        mapping,
        col_widths=[Cm(4.6), Cm(4.2), Cm(7.2)],
    )
    add_paragraph(doc, "", space_after=4)

    # ==================== 4. GIỚI HẠN & CÂU HỎI MỞ ====================
    add_h1(doc, "4. Giới hạn của dữ liệu & câu hỏi cần nghiên cứu thêm")

    add_bullet(doc, "Phát và partner là một cặp đôi trưởng thành, tự chủ tài chính ngang nhau — chưa có dữ liệu cho động lực "
                     "cha mẹ–con hoặc gia đình nhiều thế hệ, vốn là trọng tâm chính của FamilyHub",
               bold_lead="Đối tượng hẹp: ")
    add_bullet(doc, "Hộ chỉ có 2 người quản lý chi tiêu chung chặt chẽ; người thứ 3 (bạn cùng nhà) được cố tình loại khỏi hệ thống chia sẻ "
                     "chi phí (\"COD, không charge\") — cần thêm dữ liệu về hộ 3+ người chia sẻ đều",
               bold_lead="Quy mô hộ nhỏ: ")
    add_bullet(doc, "Nhu cầu chia sẻ chi tiêu ở đây gắn với một cặp đôi cụ thể, tự nguyện minh bạch với nhau vì mục tiêu chung — "
                     "chưa rõ mức độ phù hợp khi áp dụng cho gia đình có thành viên không muốn bị giám sát chi tiêu (ví dụ, tương tự "
                     "insight về Gen Z 18+ từ chối tính năng gia đình kiểu cha mẹ–con trong một phỏng vấn Momo khác)",
               bold_lead="Chưa kiểm chứng ở nhóm cưỡng chế/giám sát: ")
    add_bullet(doc, "Nên phỏng vấn thêm 2-3 hộ có cấu trúc khác (gia đình 3 thế hệ, hộ 3+ roommate chia đều, cặp cha mẹ với con đã đi làm) "
                     "để kiểm tra các đề xuất ở mục 3 có tổng quát hoá được không, trước khi đưa vào roadmap chính thức",
               bold_lead="Đề xuất tiếp theo: ")

    # ==================== 5. KẾT LUẬN ====================
    add_h1(doc, "5. Kết luận")

    add_paragraph(
        doc,
        "Điểm mấu chốt của toàn bộ phần phỏng vấn về chi tiêu: Phát không thiếu công cụ — anh đang dùng đồng thời Momo, app ngân hàng, "
        "và Google Sheets. Vấn đề là không công cụ nào trong số đó được thiết kế cho đơn vị \"hai người quản lý tiền cùng nhau nhưng vẫn "
        "giữ tài khoản riêng\", vốn là chính xác không gian sản phẩm mà FamilyHub nhắm tới. Ba khoảng trống rõ nhất — split giao dịch ở cấp "
        "bản ghi, chia sẻ tầm nhìn không cần gộp quỹ, và cảnh báo số dư khả dụng thay vì tổng chi tiêu — đều là những thứ một sản phẩm "
        "multi-member-first như FamilyHub có lợi thế tự nhiên để làm tốt hơn một ví điện tử vốn thiết kế gốc cho cá nhân rồi mới thêm "
        "tính năng chia sẻ.",
        space_after=8,
    )
    add_paragraph(
        doc,
        "Rủi ro cần lưu ý: toàn bộ phân tích này đến từ một cặp đôi hợp tác tự nguyện, không phải mọi cấu trúc gia đình đều có cùng mức độ "
        "sẵn sàng minh bạch chi tiêu với nhau. Nên xác thực thêm với các hộ có động lực khác trước khi coi các đề xuất ở mục 3 là bản thiết kế "
        "cuối cùng.",
        space_after=12,
    )

    # ==================== Footer ====================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("FamilyHub · Research Input · Nguồn: phỏng vấn MoMo (Quách Minh Phát) · 31/07/2026")
    set_run(r, size=9, italic=True, color=MUTED)

    out = "/Users/hiencao/Documents/ClaudeHC/familyhub/research/expense_management_insights_quachminhphat.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    build()
