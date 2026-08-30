#!/usr/bin/env python3
"""Render bank-email-feature-review.md into a styled .docx.

Usage: python3 docs/specs/build-bank-email-feature-review.py
Needs python-docx (pip install python-docx). Reads and writes alongside itself.
"""
import re, sys, os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = os.path.join(os.path.dirname(os.path.abspath(__file__)), "bank-email-feature-review.md")
OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "bank-email-feature-review.docx")

INK   = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x5A, 0x5A, 0x5A)
ACCENT= RGBColor(0x0B, 0x5A, 0x4A)   # deep green, matches Earthy
RULE  = RGBColor(0xC8, 0xC8, 0xC8)

doc = Document()

# ---- page + base styles -------------------------------------------------
for s in doc.sections:
    s.top_margin = Inches(0.9); s.bottom_margin = Inches(0.9)
    s.left_margin = Inches(1.0); s.right_margin = Inches(1.0)

st = doc.styles['Normal']
st.font.name = 'Georgia'; st.font.size = Pt(10.5); st.font.color.rgb = INK
st.paragraph_format.space_after = Pt(7)
st.paragraph_format.line_spacing = 1.22
rpr = st.element.get_or_add_rPr(); rf = rpr.get_or_add_rFonts()
rf.set(qn('w:eastAsia'), 'Georgia')

def style_heading(name, size, color, bold=True, before=16, after=6, caps=False):
    s = doc.styles[name]
    s.font.name = 'Helvetica Neue'; s.font.size = Pt(size)
    s.font.bold = bold; s.font.color.rgb = color
    s.font.all_caps = caps
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.keep_with_next = True

style_heading('Heading 1', 17, ACCENT, before=22, after=8)
style_heading('Heading 2', 12.5, INK,  before=16, after=5)
style_heading('Heading 3', 10.5, ACCENT, before=13, after=4)
style_heading('Heading 4', 10,   MUTED, before=11, after=3, caps=True)

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'),'clear')
    sh.set(qn('w:color'),'auto'); sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)

def hrule():
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr(); bd = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6')
    b.set(qn('w:space'),'1'); b.set(qn('w:color'),'C8C8C8')
    bd.append(b); pPr.append(bd)

INLINE = re.compile(r'(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)')

def emit_runs(par, text):
    """Bold / italic / code inline."""
    for tok in INLINE.split(text):
        if not tok: continue
        if tok.startswith('**') and tok.endswith('**') and len(tok) > 4:
            r = par.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith('`') and tok.endswith('`') and len(tok) > 2:
            r = par.add_run(tok[1:-1]); r.font.name = 'Menlo'
            r.font.size = Pt(9); r.font.color.rgb = ACCENT
        elif tok.startswith('*') and tok.endswith('*') and len(tok) > 2:
            r = par.add_run(tok[1:-1]); r.italic = True
        else:
            par.add_run(tok)

lines = open(SRC, encoding='utf-8').read().split('\n')
i = 0
first_h1 = True

while i < len(lines):
    ln = lines[i]

    # fenced code
    if ln.startswith('```'):
        i += 1; buf = []
        while i < len(lines) and not lines[i].startswith('```'):
            buf.append(lines[i]); i += 1
        i += 1
        t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
        c = t.cell(0,0); shade(c, 'F4F4F1')
        c.text = ''
        for n, b in enumerate(buf):
            p = c.paragraphs[0] if n == 0 else c.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(b if b.strip() else ' ')
            r.font.name = 'Menlo'; r.font.size = Pt(7.8)
        doc.add_paragraph().paragraph_format.space_after = Pt(3)
        continue

    # table
    if ln.startswith('|') and i+1 < len(lines) and re.match(r'^\|[\s:\-|]+\|$', lines[i+1]):
        rows = []
        while i < len(lines) and lines[i].startswith('|'):
            if not re.match(r'^\|[\s:\-|]+\|$', lines[i]):
                rows.append([c.strip() for c in lines[i].strip().strip('|').split('|')])
            i += 1
        ncol = max(len(r) for r in rows)
        t = doc.add_table(rows=0, cols=ncol); t.style = 'Table Grid'
        for ri, row in enumerate(rows):
            cells = t.add_row().cells
            for ci in range(ncol):
                txt = row[ci] if ci < len(row) else ''
                cell = cells[ci]; cell.text = ''
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                emit_runs(p, txt)
                for r in p.runs:
                    r.font.size = Pt(8.6)
                    if ri == 0: r.bold = True
                if ri == 0: shade(cell, 'EDEFEC')
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
        continue

    s = ln.strip()

    if not s:
        i += 1; continue

    if s == '---':
        hrule(); i += 1; continue

    if s.startswith('#'):
        lvl = len(s) - len(s.lstrip('#'))
        txt = s[lvl:].strip()
        if lvl == 1 and first_h1:
            p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
            r = p.add_run(txt); r.font.name = 'Helvetica Neue'
            r.font.size = Pt(24); r.bold = True; r.font.color.rgb = ACCENT
            first_h1 = False
        else:
            h = doc.add_heading(level=min(lvl, 4))
            emit_runs(h, txt)
            for r in h.runs:
                r.font.name = 'Helvetica Neue'
                r.font.color.rgb = doc.styles['Heading %d' % min(lvl,4)].font.color.rgb
                r.font.size = doc.styles['Heading %d' % min(lvl,4)].font.size
                r.bold = True
        i += 1; continue

    if s.startswith('> '):
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.32)
        p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(8)
        pPr = p._p.get_or_add_pPr(); bd = OxmlElement('w:pBdr')
        lb = OxmlElement('w:left')
        lb.set(qn('w:val'),'single'); lb.set(qn('w:sz'),'12')
        lb.set(qn('w:space'),'10'); lb.set(qn('w:color'),'0B5A4A')
        bd.append(lb); pPr.append(bd)
        emit_runs(p, s[2:])
        for r in p.runs: r.italic = True; r.font.color.rgb = MUTED
        i += 1; continue

    m = re.match(r'^(\d+)\.\s+(.*)', s)
    if m:
        p = doc.add_paragraph(style='List Number'); emit_runs(p, m.group(2))
        i += 1; continue

    if s.startswith('- '):
        p = doc.add_paragraph(style='List Bullet'); emit_runs(p, s[2:])
        i += 1; continue

    if s.startswith('*') and s.endswith('*') and not s.startswith('**'):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(s.strip('*')); r.italic = True
        r.font.size = Pt(9); r.font.color.rgb = MUTED
        i += 1; continue

    p = doc.add_paragraph(); emit_runs(p, s)
    i += 1

doc.save(OUT)
print("wrote", OUT)
