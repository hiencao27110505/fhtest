# Build effortless-transaction-logging-spec.docx from effortless-transaction-logging-spec.md
# Style: matches personal-ledger-spec.docx (Calibri, green #35654D accents,
# real Heading 1/2/3 styles so Word's navigation pane works), plus fenced
# code blocks rendered as bordered monospace paragraphs.
import re, os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = os.path.join(os.path.dirname(os.path.abspath(__file__)), "effortless-transaction-logging-spec.md")
OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "effortless-transaction-logging-spec.docx")
IMGDIR = os.path.dirname(SRC)

GREEN = RGBColor(0x35, 0x65, 0x4D)
INK = RGBColor(0x1D, 0x26, 0x21)
MUT = RGBColor(0x5C, 0x6B, 0x63)
CODE = RGBColor(0x2A, 0x4A, 0x3A)

doc = Document()

# page + default styles ------------------------------------------------------
for s in doc.sections:
    s.top_margin, s.bottom_margin = Cm(2.2), Cm(2.2)
    s.left_margin, s.right_margin = Cm(2.4), Cm(2.4)

st = doc.styles["Normal"]
st.font.name = "Calibri"; st.font.size = Pt(11); st.font.color.rgb = INK
st.paragraph_format.space_after = Pt(6); st.paragraph_format.line_spacing = 1.12

lb = doc.styles["List Bullet"]
lb.font.name = "Calibri"; lb.font.size = Pt(11); lb.font.color.rgb = INK
lb.paragraph_format.space_after = Pt(4); lb.paragraph_format.line_spacing = 1.12

def _style_heading(name, size, color, bold=True, before=18, after=6):
    h = doc.styles[name]
    h.font.name = "Calibri"; h.font.size = Pt(size); h.font.bold = bold
    h.font.color.rgb = color
    h.font.italic = False
    h.paragraph_format.space_before = Pt(before)
    h.paragraph_format.space_after = Pt(after)
    h.paragraph_format.keep_with_next = True
    return h

_style_heading("Heading 1", 19, GREEN, before=26, after=8)
_style_heading("Heading 2", 14, INK, before=20, after=6)
_style_heading("Heading 3", 11.5, GREEN, before=14, after=4)

def _add_code_run(p, text, base_size):
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt((base_size or 11) - 1)
    r.font.color.rgb = CODE
    rpr = r._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.append(rf)
    rf.set(qn("w:ascii"), "Consolas"); rf.set(qn("w:hAnsi"), "Consolas")
    return r

def _emit(p, text, bold, italic, base_size, base_color):
    """Emit runs for a bold/italic-resolved segment, honoring `code` spans."""
    for part in re.split(r"(`[^`]+?`)", text):
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            r = _add_code_run(p, part[1:-1], base_size)
            r.bold = bold
        else:
            r = p.add_run(part)
            r.bold = bold or None
            r.italic = italic or None
            if base_size:
                r.font.size = Pt(base_size)
            if base_color:
                r.font.color.rgb = base_color

def runs_from_inline(p, text, base_size=None, base_color=None):
    """Parse **bold**, *italic*, `code` into runs."""
    text = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", text)  # strip md links
    for part in re.split(r"(\*\*.+?\*\*|\*[^*]+?\*)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            _emit(p, part[2:-2], True, False, base_size, base_color)
        elif part.startswith("*") and part.endswith("*"):
            _emit(p, part[1:-1], False, True, base_size, base_color)
        else:
            _emit(p, part, False, False, base_size, base_color)
    return p

HEAD_SIZE = {1: 19, 2: 14, 3: 11.5}

def add_heading_parsed(text, level):
    p = doc.add_heading("", level=level)
    runs_from_inline(p, text, base_size=HEAD_SIZE[level])
    for r in p.runs:
        if r.font.name != "Consolas":
            r.bold = True
    return p

def add_quote_block(text):
    p = doc.add_paragraph()
    runs_from_inline(p, text, base_size=10.5, base_color=MUT)
    pf = p.paragraph_format
    pf.left_indent = Cm(0.5); pf.space_before = Pt(6); pf.space_after = Pt(8)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "10"); left.set(qn("w:color"), "35654D")
    pBdr.append(left); pPr.append(pBdr)

def add_code_block(lines_):
    """Fenced code: monospace, small, boxed with a left accent, no wrap fuss."""
    for j, ln in enumerate(lines_):
        p = doc.add_paragraph()
        _add_code_run(p, ln if ln else " ", 10)
        pf = p.paragraph_format
        pf.left_indent = Cm(0.5)
        pf.space_before = Pt(6 if j == 0 else 0)
        pf.space_after = Pt(8 if j == len(lines_) - 1 else 0)
        pf.line_spacing = 1.0
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "12")
        left.set(qn("w:space"), "10"); left.set(qn("w:color"), "C7D2CA")
        pBdr.append(left); pPr.append(pBdr)

def set_cell_borders(cell):
    tcPr = cell._element.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "C7D2CA")
        borders.append(el)
    tcPr.append(borders)

def add_table(rows):
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=ncol)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    for ri, row in enumerate(rows):
        for ci in range(ncol):
            cell_text = row[ci] if ci < len(row) else ""
            cell = t.cell(ri, ci)
            set_cell_borders(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            if ri == 0:
                clean = re.sub(r"\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`", lambda m: m.group(1) or m.group(2) or m.group(3), cell_text)
                r = p.add_run(clean)
                r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = GREEN
                tcPr = cell._element.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "EEF3EF")
                tcPr.append(shd)
            else:
                runs_from_inline(p, cell_text, base_size=9.5)
                if ci == 0:
                    for r in p.runs:
                        if not r.font.name == "Consolas":
                            r.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_image(path, alt):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.add_run().add_picture(path, width=Inches(6.1))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(alt)
    r.font.size = Pt(9); r.font.color.rgb = MUT; r.italic = True
    cap.paragraph_format.space_after = Pt(12)

# ── title block ────────────────────────────────────────────────────────────
def title_block(title, subtitle):
    p = doc.add_paragraph(); r = p.add_run("FAMILYHUB PRODUCT SPECIFICATION")
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = GREEN
    p.paragraph_format.space_after = Pt(2)
    p = doc.add_paragraph(); r = p.add_run(title)
    r.bold = True; r.font.size = Pt(24); r.font.color.rgb = INK
    p.paragraph_format.space_after = Pt(4)
    p = doc.add_paragraph(); r = p.add_run(subtitle)
    r.font.size = Pt(13); r.font.color.rgb = MUT
    p.paragraph_format.space_after = Pt(10)

# ── parse the markdown subset ──────────────────────────────────────────────
lines = open(SRC).read().split("\n")
i = 0
title_block("Effortless Transaction Logging",
            "From a bank's email to the ledger — captured, sealed, reviewed, and filed with one tap.")

# skip the md H1 title line; intro paragraphs render as normal text
while i < len(lines) and not lines[i].startswith("# "):
    i += 1
i += 1  # past '# Effortless Transaction Logging…'

para_buf = []

def flush_para():
    global para_buf
    if not para_buf:
        return
    text = " ".join(para_buf).strip()
    para_buf = []
    if not text:
        return
    # split glossary-style '**Term.** def' entries into their own paragraphs
    chunks = re.split(r"(?=\*\*[A-ZĐ“][^*]*?\.\*\*)", text) if text.startswith("**") else [text]
    for ch in chunks:
        ch = ch.strip()
        if not ch:
            continue
        p = doc.add_paragraph()
        runs_from_inline(p, ch)

while i < len(lines):
    ln = lines[i]

    if ln.startswith("```"):
        flush_para()
        code = []
        i += 1
        while i < len(lines) and not lines[i].startswith("```"):
            code.append(lines[i])
            i += 1
        add_code_block(code)
    elif ln.startswith("# "):
        flush_para()
        add_heading_parsed(ln[2:].strip(), 1)
    elif ln.startswith("## "):
        flush_para()
        add_heading_parsed(ln[3:].strip(), 2)
    elif ln.startswith("### "):
        flush_para()
        add_heading_parsed(ln[4:].strip(), 3)
    elif ln.startswith("> "):
        flush_para()
        q = [ln[2:]]
        while i + 1 < len(lines) and lines[i + 1].startswith(">"):
            i += 1
            q.append(lines[i].lstrip(">").strip())
        add_quote_block(" ".join(x for x in q if x))
    elif ln.startswith("!["):
        flush_para()
        m = re.match(r"!\[(.*?)\]\((.*?)\)", ln)
        if m:
            add_image(os.path.join(IMGDIR, m.group(2)), m.group(1))
    elif ln.startswith("|"):
        flush_para()
        tbl = []
        while i < len(lines) and lines[i].startswith("|"):
            row = [c.strip() for c in lines[i].strip("|").split("|")]
            if not all(re.fullmatch(r"[-: ]+", c) for c in row):
                tbl.append(row)
            i += 1
        i -= 1
        add_table(tbl)
    elif ln.startswith("- "):
        flush_para()
        item = [ln[2:]]
        while i + 1 < len(lines) and (lines[i + 1].startswith("  ") and lines[i + 1].strip()
                                      and not lines[i + 1].lstrip().startswith("- ")):
            i += 1
            item.append(lines[i].strip())
        p = doc.add_paragraph(style="List Bullet")
        runs_from_inline(p, " ".join(item))
    elif re.match(r"^\d+\. ", ln):
        flush_para()
        m = re.match(r"^(\d+)\. (.*)$", ln)
        item = [m.group(2)]
        while i + 1 < len(lines) and (lines[i + 1].startswith("   ") and lines[i + 1].strip()
                                      and not re.match(r"^\d+\. ", lines[i + 1].lstrip())):
            i += 1
            item.append(lines[i].strip())
        p = doc.add_paragraph(style="List Bullet")
        runs_from_inline(p, m.group(1) + ".  " + " ".join(item))
    elif ln.strip() == "---":
        flush_para()
    elif ln.strip() == "":
        flush_para()
    else:
        para_buf.append(ln.strip())
    i += 1

flush_para()

# metadata
cp = doc.core_properties
cp.title = "Effortless Transaction Logging — Product Specification"
cp.author = "FamilyHub"
cp.comments = ""
try:
    cp.last_modified_by = "FamilyHub"
except Exception:
    pass

doc.save(OUT)
print("saved", OUT)
